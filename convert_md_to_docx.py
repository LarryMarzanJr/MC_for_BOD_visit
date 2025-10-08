from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def md_to_docx(md_content, output_path):
    # Create a new document
    doc = Document()
    
    # Add title with formatting
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(md_content.split('\n')[0][2:]) # Remove '# ' from title
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Skip the first line (title) and process the rest
    lines = md_content.split('\n')[1:]
    
    current_paragraph = doc.add_paragraph()
    
    for line in lines:
        # Check for headers
        if line.startswith('## '):
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(line[3:]) # Remove '## '
            header_run.bold = True
            header_run.font.size = Pt(14)
        elif line.startswith('### '):
            sub_header_para = doc.add_paragraph()
            sub_header_run = sub_header_para.add_run(line[4:]) # Remove '### '
            sub_header_run.bold = True
            sub_header_run.font.size = Pt(12)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            clean_line = line[2:-2] # Remove '**'
            run = current_paragraph.add_run(clean_line + '\n')
            run.bold = True
        elif line.strip() == '': # Empty line
            current_paragraph = doc.add_paragraph()
        else:
            # Regular text
            run = current_paragraph.add_run(line + '\n')
    
    # Save the document
    doc.save(output_path)

# Read the markdown content
with open('/home/joenmarz/Documents/MC_for_BOD_visit/scripts/mc_mixed_id_en.md', 'r', encoding='utf-8') as file:
    md_content = file.read()

# Convert to docx
md_to_docx(md_content, '/home/joenmarz/Documents/MC_for_BOD_visit/scripts/mc_mixed_id_en.docx')

print("Successfully converted markdown to docx")